import os
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document

from io import StringIO

def read_from_pdf(file_path):
	with open(file_path, 'rb') as file:
		resource_manage = PDFResourceManager()
		return_str = StringIO()
		lap_params = LAParams()
 
		device = TextConverter(
			resource_manager, return_str, laparams=lap_params)
		process_pdf(resource_manager, device, file)
		device.close()
 
		content = return_str.getvalue()
		return_str.close()
		return content
 
def save_text_to_word(content, file_path):
	doc = Document()
	for line in content.split('\n'):
		paragraph = doc.add_paragraph()
		paragraph.add_run(remove_control_characters(line))
	doc.save(file_path)
